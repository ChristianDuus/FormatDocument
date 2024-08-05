import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import re
# Ensure dependency 'python-docx package' has been installed

# Function to create or get a style
def get_or_create_style(doc, style_name, style_type):
    try:
        return doc.styles[style_name]
    except KeyError:
        style = doc.styles.add_style(style_name, style_type)
        return style


# Function to apply styles and clean up the keys
def apply_styles(doc):
    heading1_style = get_or_create_style(doc, 'Heading 1', WD_STYLE_TYPE.PARAGRAPH)
    heading2_style = get_or_create_style(doc, 'Heading 2', WD_STYLE_TYPE.PARAGRAPH)
    list_bullet_style = get_or_create_style(doc, 'List Bullet', WD_STYLE_TYPE.PARAGRAPH)
    normal_style = get_or_create_style(doc, 'Normal', WD_STYLE_TYPE.PARAGRAPH)

    normal_style.font.size = Pt(12)
    normal_style.font.name = 'Times New Roman'

    for para in doc.paragraphs:
        # Apply styles based on the content
        if para.text.startswith('Chapter'):
            para.style = 'Heading 1'
            para.text = para.text.replace('Chapter', '').strip()
        elif para.text.startswith('### '):
            para.style = 'Heading 2'
            para.text = para.text.replace('### ', '').strip()
        elif para.text.startswith('- '):
            para.style = 'List Bullet'
            para.text = para.text.replace('- ', '').strip()
        else:
            para.style = 'Normal'

        # Center align chapter titles
        if para.style.name == 'Heading 1':
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Apply bold formatting for text enclosed in **
        if '**' in para.text:
            # Find all the text to be bolded
            bold_parts = re.findall(r'\*\*(.*?)\*\*', para.text)
            if bold_parts:
                # Split the paragraph text to process the bold parts
                parts = re.split(r'(\*\*.*?\*\*)', para.text)
                para.clear()
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        # It's a bold part
                        run = para.add_run(part[2:-2])  # Remove ** and add as bold
                        run.bold = True
                    else:
                        # It's a normal part
                        para.add_run(part)


# Load content from text file and create a Word document
txt_path = r'C:\temp\{document name}.txt'  # Use raw string for file path; update input file name
print(f"Checking if the file exists at: {txt_path}")
if not os.path.exists(txt_path):
    print(f"File not found at {txt_path}")
else:
    print("File found, proceeding to load the content.")

    try:
        # Read the content from the text file
        with open(txt_path, 'r', encoding='utf-8') as file:
            content = file.read()

        # Create a new Document
        doc = Document()

        # Add content to the document
        paragraphs = content.split('\n')
        for para in paragraphs:
            doc.add_paragraph(para)

        print("Content added to the document.")

        # Apply styles
        apply_styles(doc)

        # Save the document
        output_path = r'C:\temp\formatted_document.docx'
        doc.save(output_path)
        print(f"Document saved successfully at {output_path}")
    except Exception as e:
        print(f"An error occurred: {e}")
